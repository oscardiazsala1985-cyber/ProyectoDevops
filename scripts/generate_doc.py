"""
generate_doc.py
Genera el documento Word del reto técnico Lite Thinking 2026
Uso: python scripts/generate_doc.py
Requiere: pip install python-docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ── Configurar márgenes ───────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# ── Colores ───────────────────────────────────────────────────────────────────
COLOR_TITULO    = RGBColor(0xFF, 0x99, 0x00)   # Naranja AWS
COLOR_SECCION   = RGBColor(0x00, 0x77, 0xBB)   # Azul
COLOR_NORMAL    = RGBColor(0x1A, 0x1A, 0x2E)   # Casi negro
COLOR_GRIS      = RGBColor(0x55, 0x55, 0x55)   # Gris

def titulo_principal(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = COLOR_TITULO
    return p

def subtitulo(doc, texto):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.font.size = Pt(14)
    run.font.color.rgb = COLOR_GRIS
    return p

def heading(doc, texto, nivel=1):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(16) if nivel == 1 else Pt(13)
    run.font.color.rgb = COLOR_SECCION if nivel == 1 else COLOR_TITULO
    return p

def body(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_NORMAL
    return p

def bullet(doc, texto):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(texto)
    run.font.size = Pt(11)
    run.font.color.rgb = COLOR_NORMAL
    return p

def code_block(doc, texto):
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x88, 0x00)
    return p

def separador(doc):
    doc.add_paragraph("─" * 80)

def espacio(doc):
    doc.add_paragraph("")

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
espacio(doc)
espacio(doc)
titulo_principal(doc, "SRE Process Service")
titulo_principal(doc, "Infraestructura Cloud AWS")
espacio(doc)
subtitulo(doc, "Reto Técnico — Platform & DevOps Engineer 2026")
espacio(doc)
subtitulo(doc, "Oscar Diaz | Platform & DevOps Engineer")
subtitulo(doc, "Región: us-east-1 | IaC: Terraform | 65+ recursos AWS")
espacio(doc)
subtitulo(doc, "Repositorio: github.com/oscardiazsala1985-cyber/ProyectoDevops")
subtitulo(doc, "Servidor: https://gahk15kxi5.execute-api.us-east-1.amazonaws.com/process")
espacio(doc)
subtitulo(doc, "Junio 2026")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# RESUMEN EJECUTIVO
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "RESUMEN EJECUTIVO")
separador(doc)
body(doc, "Se implementó mediante Infraestructura como Código (IaC) con Terraform una arquitectura desacoplada y altamente disponible de 4 capas. El sistema procesa peticiones concurrentes mediante un backend elástico (EC2 y Lambda) protegido por un perímetro de red estricto (VPC privada y ALB), cuya persistencia y velocidad están respaldadas por un clúster relacional en RDS PostgreSQL y una capa de aceleración en memoria mediante Redis.")
espacio(doc)
heading(doc, "Arquitectura de 4 Capas", nivel=2)
bullet(doc, "Capa 1 — Ingesta Dual: API Gateway (serverless) + ALB (servidores EC2)")
bullet(doc, "Capa 2 — Cómputo Híbrido: Lambda + Auto Scaling Group (EC2 t3.micro)")
bullet(doc, "Capa 3 — Datos y Caché: RDS PostgreSQL 16.3 + ElastiCache Redis 7.1 + S3")
bullet(doc, "Capa 4 — Gobernanza: CloudTrail + GuardDuty + IAM + CloudWatch + X-Ray + Grafana")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PUNTO 1 — ADMINISTRACIÓN DE INFRAESTRUCTURA
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "PUNTO 1 — Administración de Infraestructura")
separador(doc)

heading(doc, "¿Qué se implementó?", nivel=2)
body(doc, "Se desplegaron instancias EC2 (Amazon Linux 2023, t3.micro) dentro de un Auto Scaling Group en subnets privadas. Las instancias no tienen IP pública y el acceso administrativo se realiza exclusivamente via AWS Systems Manager Session Manager (SSM), eliminando la necesidad de SSH y llaves de acceso.")
espacio(doc)

heading(doc, "Alta Disponibilidad implementada:", nivel=2)
bullet(doc, "ALB multi-AZ en 2 subnets públicas (AZ-a y AZ-b) — si una AZ cae, el tráfico va a la otra")
bullet(doc, "ASG: min=1 / desired=2 / max=4 — escala automáticamente por CPU al 60%")
bullet(doc, "Health checks ELB cada 30s en /health — instancias unhealthy se reemplazan sin intervención")
bullet(doc, "Rolling Update con min_healthy_percentage=50 — actualizaciones sin downtime")
espacio(doc)

heading(doc, "Actualización de SO sin afectar el servicio:", nivel=2)
code_block(doc, "aws autoscaling start-instance-refresh \\")
code_block(doc, "  --auto-scaling-group-name sre-process-service-dev-asg \\")
code_block(doc, "  --preferences '{\"MinHealthyPercentage\": 50}'")
body(doc, "El ASG reemplaza instancias de a una: drena conexiones → termina instancia vieja → lanza nueva con AMI actualizada → espera health check → continúa con la siguiente.")
espacio(doc)

heading(doc, "Entorno mixto Linux/Windows:", nivel=2)
bullet(doc, "Active Directory: AWS Managed Microsoft AD + ssm send-command para unir instancias al dominio")
bullet(doc, "Group Policies: aplicadas automáticamente desde el AD al reiniciar instancias")
bullet(doc, "PowerShell via SSM Run Command — sin necesidad de RDP ni credenciales expuestas")
body(doc, "NOTA: Ver capturas de VMs locales en el Punto 3 (Ubuntu Server + Windows Server 2025)")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PUNTO 2 — SERVICIOS DE RED
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "PUNTO 2 — Servicios de Red")
separador(doc)

heading(doc, "Configuración de red AWS desplegada:", nivel=2)
bullet(doc, "VPC: 10.40.0.0/16 con DNS hostnames habilitado")
bullet(doc, "Subnets públicas: 10.40.0.0/24 (us-east-1a) y 10.40.1.0/24 (us-east-1b)")
bullet(doc, "Subnets privadas: 10.40.10.0/24 (us-east-1a) y 10.40.11.0/24 (us-east-1b)")
bullet(doc, "Internet Gateway → subnets públicas | NAT Gateway → subnets privadas")
bullet(doc, "VPC Endpoint S3 Gateway — tráfico a S3 sin salir a Internet (gratis + más rápido)")
espacio(doc)

heading(doc, "Security Groups (Zero Trust entre capas):", nivel=2)
bullet(doc, "ALB-SG: ingress 80/443 Internet | egress 8080 → EC2-SG")
bullet(doc, "EC2-SG: ingress 8080 ALB-SG | egress 6379 Redis-SG, 5432 RDS-SG")
bullet(doc, "Lambda-SG: egress 6379 Redis-SG, 5432 RDS-SG, 443 Internet")
bullet(doc, "Redis-SG: ingress 6379 Lambda-SG, EC2-SG — sin acceso público")
bullet(doc, "RDS-SG: ingress 5432 EC2-SG, Lambda-SG — sin acceso público")
espacio(doc)

heading(doc, "Decisiones de segmentación:", nivel=2)
bullet(doc, "ALB y NAT en subnets públicas — los servidores nunca exponen IPs públicas")
bullet(doc, "RDS, Redis, Lambda y EC2 en subnets privadas — aislamiento total")
bullet(doc, "Security Groups por SG-source (no por CIDR) — más seguro y flexible")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PUNTO 3 — CONTENEDORES Y VIRTUALIZACIÓN
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "PUNTO 3 — Contenedores y Virtualización")
separador(doc)

heading(doc, "VMs locales desplegadas:", nivel=2)
body(doc, "Se configuraron y gestionaron dos máquinas virtuales en entorno local:")
espacio(doc)
bullet(doc, "VM 1 — Ubuntu Server 22.04 LTS en Oracle VirtualBox")
code_block(doc, "  Hostname: oscardiaz | Virtualización: oracle")
code_block(doc, "  OS: Ubuntu 22.04.5 LTS | Kernel: 5.15.0-157-generic")
code_block(doc, "  Docker instalado (interfaz docker0 visible)")
code_block(doc, "  IP: 192.168.1.94 | Conectividad verificada")
espacio(doc)
bullet(doc, "VM 2 — Windows Server 2025 en VMware Workstation")
code_block(doc, "  Windows 10.0.26100 | Hostname: oscardiaz")
code_block(doc, "  Ping 8.8.8.8: 4/4 paquetes, 0 pérdidas, latencia ~49ms")
code_block(doc, "  3 unidades: C: (Windows), A: (datos), D: (DVD)")
espacio(doc)

heading(doc, "Almacenamiento persistente en contenedores:", nivel=2)
bullet(doc, "Docker Volumes: datos persisten entre reinicios (redis_data en docker-compose.yml)")
bullet(doc, "Kubernetes PVC con EBS gp3 (ReadWriteOnce) — para Redis en EKS")
bullet(doc, "Kubernetes PVC con EFS (ReadWriteMany) — para logs compartidos entre pods")
bullet(doc, "En producción: RDS PostgreSQL (managed), ElastiCache Redis (snapshots), S3 (objetos)")
espacio(doc)

heading(doc, "Docker implementado:", nivel=2)
bullet(doc, "Dockerfile multi-stage con usuario no-root y healthcheck")
bullet(doc, "docker-compose.yml: App + Redis + Redis Commander UI en un comando")
bullet(doc, "k8s/: 7 manifiestos Kubernetes (Deployment, Service, HPA, PVC, ConfigMap, Secret)")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PUNTO 4 — CIBERSEGURIDAD
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "PUNTO 4 — Ciberseguridad Integrada")
separador(doc)

heading(doc, "Prácticas implementadas:", nivel=2)
bullet(doc, "Security Groups desacoplados — reglas por SG-source, no por CIDR")
bullet(doc, "IMDSv2 obligatorio en EC2 — previene ataques SSRF al metadata service")
bullet(doc, "Cifrado en reposo: RDS (AES-256), S3 (AES-256), EBS (gp3 cifrado)")
bullet(doc, "Solo HTTPS — S3 bucket policy niega conexiones HTTP")
bullet(doc, "Sin credenciales en código — Lambda e EC2 usan IAM Roles")
bullet(doc, "db_password via TF_VAR — nunca hardcodeado en git")
espacio(doc)

heading(doc, "Amazon GuardDuty — Detección de amenazas con ML:", nivel=2)
body(doc, "GuardDuty analiza CloudTrail, VPC Flow Logs y DNS logs usando machine learning para detectar comportamiento anómalo sin instalar agentes.")
bullet(doc, "Malware Protection en volúmenes EBS habilitado")
bullet(doc, "S3 data events monitoreados")
bullet(doc, "Findings de severidad Media/Alta/Crítica → SNS en 15 minutos")
bullet(doc, "Resultado: 1 finding detectado (uso de root credentials) — GuardDuty funcionando")
espacio(doc)

heading(doc, "AWS CloudTrail activo:", nivel=2)
bullet(doc, "Trail multi-región — captura eventos en TODAS las regiones")
bullet(doc, "Log file validation SHA-256 — detecta tampering de logs")
bullet(doc, "Logs en S3 (90 días) + CloudWatch Logs (tiempo real)")
bullet(doc, "Alarmas: cambios IAM detectados, AccessDenied repetidos")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PUNTO 5 — GESTIÓN DE NUBE E IAM
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "PUNTO 5 — Gestión de Nube e IAM")
separador(doc)

heading(doc, "RDS PostgreSQL desplegado:", nivel=2)
bullet(doc, "Motor: PostgreSQL 16.3 | Instancia: db.t3.micro | Storage: 20GB→100GB gp3")
bullet(doc, "Cifrado AES-256 en reposo | publicly_accessible = false")
bullet(doc, "Backup 7 días con ventana 02:00-03:00 | PITR habilitado (RPO 5 min)")
bullet(doc, "Performance Insights 7 días | Enhanced Monitoring cada 60s")
bullet(doc, "Logs PostgreSQL → CloudWatch automáticamente")
espacio(doc)

heading(doc, "IAM — Principio de Mínimo Privilegio:", nivel=2)
bullet(doc, "lambda-exec-role: S3 bucket propio + VPC + CloudWatch + X-Ray SOLO")
bullet(doc, "ec2-app-role: SSM + CloudWatch Agent + X-Ray SOLO")
bullet(doc, "rds-monitoring-role: SOLO métricas RDS Enhanced Monitoring")
bullet(doc, "cloudtrail-cw-role: SOLO CreateLogStream + PutLogEvents en su log group")
bullet(doc, "grafana-role: SOLO lectura de métricas y logs CloudWatch")
espacio(doc)

heading(doc, "Política MFA para operaciones destructivas:", nivel=2)
code_block(doc, '{"Effect":"Deny","Action":["rds:DeleteDBInstance","ec2:TerminateInstances"],')
code_block(doc, ' "Condition":{"BoolIfExists":{"aws:MultiFactorAuthPresent":"false"}}}')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PUNTO 6 — SOPORTE 24/7
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "PUNTO 6 — Soporte 24/7 y Postmortem ITIL")
separador(doc)

heading(doc, "Proceso de respuesta ante incidente:", nivel=2)
bullet(doc, "1. DETECCIÓN → CloudWatch Alarm → SNS → PagerDuty/SMS")
bullet(doc, "2. ACKNOWLEDGE → Confirmar en < 5 minutos (SLA P1)")
bullet(doc, "3. TRIAGE → P1/P2/P3 según impacto")
bullet(doc, "4. COMUNICAR → Canal Slack + notificar stakeholders ETA 30min")
bullet(doc, "5. DIAGNOSTICAR → CloudWatch Logs + X-Ray + métricas ALB/ASG")
bullet(doc, "6. MITIGAR → Rollback, reinicio o failover")
bullet(doc, "7. POSTMORTEM → Documento formal en 48h")
espacio(doc)

heading(doc, "Ejemplo de incidente real:", nivel=2)
body(doc, "02:30 AM — Lambda deja de responder. API Gateway devuelve 502.")
code_block(doc, "02:30 → Alarma lambda-errors > 5 → SNS → PagerDuty")
code_block(doc, "02:37 → X-Ray muestra timeout en Redis (3000ms exceeded)")
code_block(doc, "02:42 → ElastiCache auto-recuperado. Servicio restaurado.")
code_block(doc, "02:50 → Incidente cerrado. Duración: 20 min. SLA: 99.94%")
espacio(doc)

heading(doc, "Postmortem ITIL:", nivel=2)
bullet(doc, "Causa raíz: ventana de mantenimiento ElastiCache no ajustada a horario de verano")
bullet(doc, "Impacto: ~240 peticiones fallidas. SLA mensual 99.94% (objetivo 99.9% ✓)")
bullet(doc, "Correctiva: actualizar maintenance_window + retry con backoff en Lambda")
bullet(doc, "Preventiva: circuit breaker para Redis + ElastiCache Serverless en evaluación")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PUNTO 7 — DISEÑO DE INFRAESTRUCTURA
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "PUNTO 7 — Diseño de Infraestructura Microservicios")
separador(doc)

heading(doc, "Cloud PaaS — AWS EKS:", nivel=2)
body(doc, "Internet → CloudFront → ALB → EKS Cluster (Deployment 2 réplicas + HPA 2→10 pods + Service LoadBalancer + PVC EBS/EFS) → Managed: RDS + ElastiCache + S3 + ECR")
espacio(doc)

heading(doc, "On-Premise — kubeadm + VMware:", nivel=2)
body(doc, "Firewall → HAProxy/MetalLB → Kubernetes kubeadm (Master Nodes x3 HA etcd + Worker Nodes VMware VMs) → Storage: Longhorn/Ceph → Registry: Harbor → Monitoreo: Prometheus + Grafana")
espacio(doc)

heading(doc, "Cloud IaaS — EC2 self-managed (arquitectura actual):", nivel=2)
body(doc, "ALB → EC2 ASG (Docker Engine + app Node.js :8080) → EBS Volumes (PersistentVolumes) → EFS (logs compartidos) → RDS PostgreSQL + ElastiCache Redis + S3")
espacio(doc)

heading(doc, "Manifiestos Kubernetes implementados (k8s/):", nivel=2)
bullet(doc, "namespace.yaml — aislamiento de recursos")
bullet(doc, "deployment.yaml — 2 réplicas, rolling update, anti-affinity, liveness/readiness probes")
bullet(doc, "service.yaml — LoadBalancer app + ClusterIP Redis")
bullet(doc, "hpa.yaml — autoscaler 2→10 pods por CPU 60% y RAM 70%")
bullet(doc, "pvc.yaml — EBS gp3 (Redis) + EFS (logs compartidos)")
bullet(doc, "configmap.yaml + secret.yaml — variables seguras")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PUNTO 8 — RESPALDO Y RECUPERACIÓN
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "PUNTO 8 — Políticas de Respaldo y Recuperación")
separador(doc)

heading(doc, "AWS Backup desplegado (backup.tf):", nivel=2)
bullet(doc, "Vault cifrado AES-256: sre-process-service-dev-backup-vault")
bullet(doc, "Plan diario: 02:00 UTC → retención 14 días (RPO 24h, RTO ~15 min)")
bullet(doc, "Plan semanal: Domingos 03:00 UTC → retención 120 días (compliance)")
bullet(doc, "RDS PITR continuo: RPO 5 minutos")
bullet(doc, "Alarma CloudWatch si job de backup falla")
espacio(doc)

heading(doc, "Prueba automatizada de restore (scripts/backup_test.sh):", nivel=2)
body(doc, "Script Bash que ejecuta mensualmente:")
bullet(doc, "1. Encuentra el snapshot más reciente de RDS")
bullet(doc, "2. Restaura en instancia temporal db.t3.micro")
bullet(doc, "3. Valida integridad de la base de datos")
bullet(doc, "4. Elimina instancia temporal (evita costos)")
bullet(doc, "5. Guarda resultado en backup_test_history.csv")
espacio(doc)

heading(doc, "RPO y RTO por servicio:", nivel=2)
bullet(doc, "RDS (PITR): RPO 5 min | RTO ~30 min")
bullet(doc, "RDS (snapshot diario): RPO 24h | RTO ~15 min")
bullet(doc, "S3 (versioning): RPO 0 | RTO inmediato")
bullet(doc, "ElastiCache: RPO 24h | RTO ~15 min")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PUNTO 9 — IaC Y AUTOMATIZACIÓN
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "PUNTO 9 — IaC y Automatización con Terraform")
separador(doc)

heading(doc, "Terraform implementado:", nivel=2)
bullet(doc, "16 archivos .tf | 65+ recursos AWS | Terraform >= 1.6.0 | Provider AWS ~5.0")
bullet(doc, "Plan: terraform plan → muestra cambios SIN aplicar (sin sorpresas)")
bullet(doc, "Apply: terraform apply → despliega toda la infraestructura en ~15 min")
bullet(doc, "Destroy: terraform destroy → elimina todo limpiamente (gestión de costos)")
espacio(doc)

heading(doc, "Archivos Terraform:", nivel=2)
bullet(doc, "network.tf — VPC, subnets, IGW, NAT, SGs, VPC Endpoint S3")
bullet(doc, "ec2_compute.tf — ALB, ASG, Launch Template, scaling policy")
bullet(doc, "lambda.tf — función, X-Ray, variables de entorno DB+Redis")
bullet(doc, "rds.tf — PostgreSQL, subnet group, monitoring role")
bullet(doc, "redis.tf — ElastiCache Redis 7.1")
bullet(doc, "cloudtrail.tf — trail multi-región, S3 audit, alarmas IAM")
bullet(doc, "guardduty.tf — detector ML, SNS alerts, CW alarm")
bullet(doc, "backup.tf — vault, plan diario/semanal, alarma fallo")
bullet(doc, "monitoring.tf — Grafana EC2, CloudWatch Dashboard 9 widgets")
espacio(doc)

heading(doc, "Scripts de automatización:", nivel=2)
bullet(doc, "scripts/backup_test.sh — prueba automatizada restore RDS")
bullet(doc, "scripts/health_check.sh — 5 smoke tests del endpoint")
bullet(doc, "scripts/deploy_local.sh — start/stop/restart entorno Docker")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PUNTO 10 — MONITOREO
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "PUNTO 10 — Monitoreo y Optimización de Rendimiento")
separador(doc)

heading(doc, "Herramientas implementadas:", nivel=2)
bullet(doc, "Grafana v13.1.0 — Dashboard visual conectado a CloudWatch via IAM Role")
bullet(doc, "AWS CloudWatch — 5 Log Groups + 7 Alarmas + Dashboard 9 widgets")
bullet(doc, "AWS X-Ray — Active Tracing 100% invocaciones Lambda")
bullet(doc, "Amazon GuardDuty — detección de amenazas ML 24/7")
bullet(doc, "AWS CloudTrail — auditoría completa de todas las llamadas API")
espacio(doc)

heading(doc, "Métricas monitoreadas en Grafana + CloudWatch:", nivel=2)
bullet(doc, "Lambda: invocaciones, errores, duración — 4 invocaciones, 0 errores, 100% éxito")
bullet(doc, "EC2 ASG: CPU utilization — target 60%, alarma a 80%")
bullet(doc, "ALB: peticiones, errores 5xx, latencia p50/p95/p99")
bullet(doc, "RDS: CPU, conexiones, storage libre, latencia de queries")
bullet(doc, "Redis: Cache HITs vs MISSes — demostrado con X-Cache: HIT")
espacio(doc)

heading(doc, "Auto-scaling sin intervención manual:", nivel=2)
body(doc, "Target Tracking Scaling sobre CPU 60%. Si CPU sube: agrega instancia cada 60s. Si CPU baja: reduce con ventana de 5 minutos para evitar oscilaciones.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PUNTO 11 — CI/CD
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "PUNTO 11 — Buenas Prácticas de CI/CD")
separador(doc)

heading(doc, "CI vs CD:", nivel=2)
bullet(doc, "CI (Integración Continua): corre en cada commit → build + tests + lint + security scan")
bullet(doc, "CD (Entrega Continua): corre al mergear a main → deploy DEV auto + PROD con gate manual")
espacio(doc)

heading(doc, "Pipeline implementado (.github/workflows/):", nivel=2)
bullet(doc, "ci.yml: Lint + Tests Jest (7 unitarios + 2 integración) + Docker build + Terraform validate + tfsec")
bullet(doc, "deploy.yml: DEV automático + smoke test | PROD con aprobación manual + rollback si falla")
bullet(doc, "Notificación de fallo: job notify-failure reporta commit, rama, actor y enlace al run")
espacio(doc)

heading(doc, "Tests automatizados (lambda/index.test.js):", nivel=2)
bullet(doc, "HTTP 200 con payload válido")
bullet(doc, "Campos id, processedAt, algorithm presentes en respuesta")
bullet(doc, "X-Cache: MISS en primera petición")
bullet(doc, "Flujo completo MISS → HIT con mismo payload")
bullet(doc, "Payloads distintos generan IDs distintos")
bullet(doc, "Manejo de body vacío sin crashear")
bullet(doc, "Manejo de body base64 correctamente")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# PUNTO 12 — DOCUMENTACIÓN
# ══════════════════════════════════════════════════════════════════════════════
heading(doc, "PUNTO 12 — Documentación y Mejora Continua")
separador(doc)

heading(doc, "README.md — 1300+ líneas cubriendo:", nivel=2)
bullet(doc, "Instrucciones de despliegue paso a paso (8 pasos desde git clone hasta curl)")
bullet(doc, "Decisiones de arquitectura con justificación técnica")
bullet(doc, "Referencias a cada punto del reto")
bullet(doc, "Tabla de variables, outputs y recursos")
bullet(doc, "Respuestas técnicas a los 12 puntos del reto")
espacio(doc)

heading(doc, "¿Por qué documentar cada cambio?", nivel=2)
bullet(doc, "Trazabilidad: git log muestra quién cambió qué y cuándo")
bullet(doc, "Reproducibilidad: cualquier ingeniero puede recrear el entorno desde cero")
bullet(doc, "Compliance: auditorías requieren evidencia de cambios y aprobaciones")
bullet(doc, "Reducción de MTTR: documentación reduce tiempo de diagnóstico en incidentes")
bullet(doc, "Onboarding: nuevo miembro productivo en horas, no semanas")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# CIERRE
# ══════════════════════════════════════════════════════════════════════════════
espacio(doc)
espacio(doc)
titulo_principal(doc, "Entregables")
separador(doc)
espacio(doc)
heading(doc, "Repositorio GitHub:", nivel=2)
body(doc, "https://github.com/oscardiazsala1985-cyber/ProyectoDevops")
espacio(doc)
heading(doc, "Servidor accesible (API Gateway):", nivel=2)
body(doc, "https://gahk15kxi5.execute-api.us-east-1.amazonaws.com/process")
espacio(doc)
heading(doc, "Servidor accesible (Grafana):", nivel=2)
body(doc, "http://sre-process-service-dev-alb-148576592.us-east-1.elb.amazonaws.com/grafana")
espacio(doc)
heading(doc, "Prueba del endpoint:", nivel=2)
code_block(doc, 'curl -X POST https://gahk15kxi5.execute-api.us-east-1.amazonaws.com/process \\')
code_block(doc, '  -H "Content-Type: application/json" \\')
code_block(doc, '  -d \'{"data": "lite-thinking-2026"}\'')
espacio(doc)
espacio(doc)
titulo_principal(doc, "Oscar Diaz")
subtitulo(doc, "Platform & DevOps Engineer")
subtitulo(doc, "Junio 2026")

# ── Guardar ───────────────────────────────────────────────────────────────────
output_path = "Reto_Tecnico_DevOps_Oscar_Diaz.docx"
doc.save(output_path)
print(f"✅ Documento generado: {output_path}")
